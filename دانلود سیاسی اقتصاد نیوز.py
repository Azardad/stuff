import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from io import BytesIO
from urllib.parse import urljoin
import os
import threading
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Flags for controlling the script
stop_script = False
condition = threading.Condition()

def read_last_downloaded_article(file_path='last_downloaded_article.txt'):
    if not os.path.exists(file_path):
        return None
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read().strip()

def write_last_downloaded_article(url, file_path='last_downloaded_article.txt'):
    with open(file_path, 'w', encoding='utf-8') as file:
        file.write(url)

def extract_article_links(page_url):
    try:
        response = requests.get(page_url)
        soup = BeautifulSoup(response.content, 'html.parser')

        article_links = []
        articles_div = soup.find('div', class_='category-main-content')
        if articles_div:
            right_div = articles_div.find('div', class_='right')
            if right_div:
                for a_tag in right_div.find_all('a', href=True):
                    article_url = a_tag['href']
                    if article_url.startswith('/'):
                        article_url = urljoin(page_url, article_url)
                    if 'page=' not in article_url:
                        article_links.append(article_url)

        return article_links

    except Exception:
        pass  # Silencing the exception messages

def reverse_parentheses(text):
    # Replaces '(' with ')' and ')' with '('
    return text.replace('(', 'TEMP').replace(')', '(').replace('TEMP', ')')

def set_rtl_paragraph(paragraph):
    paragraph.alignment = 3  # Align Right
    paragraph_format = paragraph.paragraph_format
    paragraph_format.right_to_left = True  # Set text direction to RTL

    # Full RTL support in Word
    bidi_element = OxmlElement('w:bidi')
    bidi_element.set(qn('w:val'), "1")
    paragraph._element.get_or_add_pPr().append(bidi_element)

def process_article(url):
    try:
        response = requests.get(url)
        soup = BeautifulSoup(response.content, 'html.parser')

        article_tag = soup.find('article')

        if not article_tag:
            return

        header_tag = article_tag.find('header')
        if header_tag:
            title = header_tag.find('h1').get_text(strip=True) if header_tag.find('h1') else "عنوان یافت نشد"
            date = header_tag.find('time').get_text(strip=True) if header_tag.find('time') else "تاریخ یافت نشد"
        else:
            title = "عنوان یافت نشد"
            date = "تاریخ یافت نشد"

        doc = Document()

        title_paragraph = doc.add_heading(title, level=1)
        set_rtl_paragraph(title_paragraph)

        date_paragraph = doc.add_paragraph(f" {date}")

        # Add article content
        for element in article_tag.find_all(['p', 'img']):
            if element.name == 'p':
                para_text = reverse_parentheses(element.get_text(strip=True))
                para = doc.add_paragraph(para_text)
                set_rtl_paragraph(para)
            elif element.name == 'img':
                img_url = element.get('src')
                if not img_url.startswith('http'):
                    img_url = urljoin(url, img_url)
                try:
                    img_data = requests.get(img_url).content
                    img_stream = BytesIO(img_data)
                    doc.add_picture(img_stream, width=Inches(4))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = 1  # Center align the image
                except Exception:
                    pass  # Silencing the exception messages

        # Add mn-header and mn-body content
        header_div = soup.find('div', class_='mn-header')
        body_div = soup.find('div', class_='mn-body')

        if header_div:
            header_paragraph = doc.add_paragraph(header_div.get_text(strip=True))
            set_rtl_paragraph(header_paragraph)

        if body_div:
            for li in body_div.find_all('li'):
                link_text = li.get_text(strip=True)
                if link_text:
                    para = doc.add_paragraph(link_text)
                    set_rtl_paragraph(para)

        # Define the folder where you want to save the documents
       # save_folder = "hello"

        # Ensure the folder exists
        #os.makedirs(save_folder, exist_ok=True)

        # Save the document with a safe title in the specified folder
        safe_title = title.replace("ـ", " ").replace("-", " ").replace("/", "_").replace(":", "_")
        doc_filename = os.path.join(f"{safe_title}.docx")
        doc.save(doc_filename)

    except Exception:
        pass  # Silencing the exception messages

def stop_program():
    global stop_script
    input("Press Enter to stop the program...")  # Instructions in English
    stop_script = True
    os._exit(0)  # Close the terminal after stopping

# Start the stop program thread
stop_thread = threading.Thread(target=stop_program, daemon=True)
stop_thread.start()

base_url = 'https://www.eghtesadnews.com/%D8%A8%D8%AE%D8%B4-%D8%A7%D8%AE%D8%A8%D8%A7%D8%B1-%D8%B3%DB%8C%D8%A7%D8%B3%DB%8C-57?page=1'
page_number = 1

# Read the last downloaded article
last_downloaded_article = read_last_downloaded_article()

first_new_article_saved = False  # Flag to check if the first new article has been saved

while not stop_script:
    page_url = f"{base_url}&page={page_number}"
    article_links = extract_article_links(page_url)

    if not article_links:
        break

    new_articles_found = False

    for link in article_links:
        if stop_script:
            break
        if link == last_downloaded_article:
            print("Reached the last downloaded article. Exiting.")
            os._exit(0)  # Close the terminal automatically

        # Only save the first new article
        if not first_new_article_saved:
            write_last_downloaded_article(link)
            first_new_article_saved = True

        process_article(link)
        new_articles_found = True

    if not new_articles_found:
        break

    page_number += 1

print("No new articles found. Exiting.")
os._exit(0)  # Close the terminal automatically
